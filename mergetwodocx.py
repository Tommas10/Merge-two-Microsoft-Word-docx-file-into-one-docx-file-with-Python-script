#!/usr/bin/env python

# This is small automation to merge two Microsoft Word docx file with Python. 
# Created by Tommas Huang.

#python-docx is a Python library for creating and updating Microsoft Word (.docx) files.
from docx import Document

# Merge Microsoft Word files source.
t1 = Document("/Users/TommasHuang/Documents/test1.docx")
# # Merge Microsoft Word files distance.
t2 = Document("/Users/TommasHuang/Documents/test2.docx")
for p in t2.paragraphs:
    # Here is the new paragraph, followed by the style is the style
    t1.add_paragraph(p.text,p.style)
# Merge Microsoft Word two files save new file
t1.save("test1-new.docx")
